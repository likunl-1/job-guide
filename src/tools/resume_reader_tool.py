"""
简历文件读取工具
支持读取Word（.docx）和PDF（.pdf）格式的简历文件
"""

from langchain.tools import tool, ToolRuntime
import os
from typing import Optional


def _read_word_docx(file_path: str) -> str:
    """
    读取Word文档（.docx格式）
    
    Args:
        file_path: Word文件路径
        
    Returns:
        文件文本内容
        
    Raises:
        ImportError: 如果未安装python-docx
        FileNotFoundError: 如果文件不存在
        Exception: 其他读取错误
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("未安装python-docx库，请运行：pip install python-docx")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在：{file_path}")
    
    doc = Document(file_path)
    
    # 提取所有段落的文本
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    
    # 提取表格中的文本
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                tables_text.append(" | ".join(row_text))
    
    # 合并段落和表格内容
    full_text = "\n".join(paragraphs)
    if tables_text:
        full_text += "\n\n【表格内容】\n" + "\n".join(tables_text)
    
    return full_text


def _read_pdf(file_path: str) -> str:
    """
    读取PDF文档（.pdf格式）
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        文件文本内容
        
    Raises:
        ImportError: 如果未安装pdfplumber
        FileNotFoundError: 如果文件不存在
        Exception: 其他读取错误
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("未安装pdfplumber库，请运行：pip install pdfplumber")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在：{file_path}")
    
    text_lines = []
    
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_lines.append(f"--- 第 {page_num} 页 ---")
                text_lines.append(page_text)
    
    return "\n".join(text_lines)


def _read_text_file(file_path: str) -> str:
    """
    读取文本文件（.txt, .md等）
    
    Args:
        file_path: 文本文件路径
        
    Returns:
        文件文本内容
        
    Raises:
        FileNotFoundError: 如果文件不存在
        Exception: 读取错误
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在：{file_path}")
    
    # 尝试多种编码
    encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    
    # 如果所有编码都失败，使用二进制模式读取
    with open(file_path, 'rb') as f:
        return f.read().decode('utf-8', errors='ignore')


@tool
def read_resume_file(
    file_path: str, 
    runtime: ToolRuntime = None
) -> str:
    """
    读取简历文件（支持Word、PDF、文本格式）
    
    Args:
        file_path: 简历文件路径，支持以下格式：
            - Word文档：.docx
            - PDF文档：.pdf
            - 文本文件：.txt, .md
            文件路径可以是：
            - 相对路径（相对于项目根目录）：resumes/my_resume.docx
            - 绝对路径：/path/to/resume.pdf
            - assets目录下：assets/resumes/my_resume.pdf
        
        runtime: LangChain工具运行时
        
    Returns:
        简历文件的文本内容
        
    Examples:
        >>> read_resume_file("assets/resumes/my_resume.pdf")
        "个人简历\n姓名：张三\n..."
        
        >>> read_resume_file("resumes/resume.docx")
        "个人简历\n姓名：李四\n..."
    """
    workspace_path = os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects")
    
    # 处理相对路径
    if not os.path.isabs(file_path):
        # 如果路径不以/开头，先尝试相对于workspace
        absolute_path = os.path.join(workspace_path, file_path)
    else:
        absolute_path = file_path
    
    # 检查文件是否存在
    if not os.path.exists(absolute_path):
        return f"❌ 错误：文件不存在\n\n文件路径：{absolute_path}\n\n请确认：\n1. 文件路径是否正确\n2. 文件是否已上传到指定目录\n3. 检查文件名和扩展名"
    
    # 获取文件扩展名
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # 根据扩展名选择读取方法
    try:
        if file_ext == '.docx':
            content = _read_word_docx(absolute_path)
            file_type = "Word文档"
        elif file_ext == '.pdf':
            content = _read_pdf(absolute_path)
            file_type = "PDF文档"
        elif file_ext in ['.txt', '.md']:
            content = _read_text_file(absolute_path)
            file_type = "文本文件"
        else:
            return f"❌ 错误：不支持的文件格式\n\n文件路径：{file_path}\n文件类型：.{file_ext}\n\n支持的格式：\n- Word文档：.docx\n- PDF文档：.pdf\n- 文本文件：.txt, .md"
        
        # 检查内容是否为空
        if not content.strip():
            return f"⚠️ 警告：文件内容为空\n\n文件路径：{file_path}\n文件类型：{file_type}\n\n可能原因：\n1. 文件是空白的\n2. 文件是扫描版PDF（图片格式）\n3. 文件编码有问题"
        
        # 返回成功读取的内容
        result = f"✅ 成功读取简历文件\n\n"
        result += f"**文件路径**：{file_path}\n"
        result += f"**文件类型**：{file_type}\n"
        result += f"**内容长度**：{len(content)} 字符\n\n"
        result += "--- 简历内容开始 ---\n\n"
        result += content
        result += "\n\n--- 简历内容结束 ---"
        
        return result
        
    except ImportError as e:
        return f"❌ 缺少必要的库\n\n{str(e)}\n\n请安装：\n- Word文档：pip install python-docx\n- PDF文档：pip install pdfplumber"
        
    except FileNotFoundError as e:
        return f"❌ 文件不存在\n\n{str(e)}\n\n请确认：\n1. 文件路径是否正确\n2. 文件是否已上传"
        
    except Exception as e:
        return f"❌ 读取文件时出错\n\n错误信息：{str(e)}\n\n文件路径：{file_path}\n\n建议：\n1. 检查文件是否损坏\n2. 尝试使用其他格式\n3. 联系技术支持"


@tool
def list_resume_files(
    directory: str = "assets/resumes",
    runtime: ToolRuntime = None
) -> str:
    """
    列出指定目录下的简历文件
    
    Args:
        directory: 简历目录路径（默认：assets/resumes）
        runtime: LangChain工具运行时
        
    Returns:
        目录下所有简历文件的列表
        
    Examples:
        >>> list_resume_files()
        "✅ 找到3个简历文件：\n\n1. 张三_简历.pdf\n2. 李四_简历.docx\n3. 王五_简历.txt"
        
        >>> list_resume_files("my_resumes")
        "✅ 找到2个简历文件：\n\n..."
    """
    workspace_path = os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects")
    
    # 处理相对路径
    if not os.path.isabs(directory):
        absolute_dir = os.path.join(workspace_path, directory)
    else:
        absolute_dir = directory
    
    # 检查目录是否存在
    if not os.path.exists(absolute_dir):
        return f"❌ 目录不存在\n\n目录路径：{absolute_dir}\n\n建议：\n1. 创建目录：mkdir -p {directory}\n2. 将简历文件放入该目录"
    
    # 支持的文件扩展名
    supported_extensions = ['.pdf', '.docx', '.txt', '.md']
    
    # 查找所有简历文件
    resume_files = []
    for filename in os.listdir(absolute_dir):
        file_ext = os.path.splitext(filename)[1].lower()
        if file_ext in supported_extensions:
            file_path = os.path.join(absolute_dir, filename)
            file_size = os.path.getsize(file_path)
            resume_files.append({
                'name': filename,
                'path': os.path.join(directory, filename),
                'type': file_ext,
                'size': file_size
            })
    
    # 按文件名排序
    resume_files.sort(key=lambda x: x['name'])
    
    # 生成返回结果
    if not resume_files:
        return f"⚠️ 目录下没有找到简历文件\n\n目录路径：{absolute_dir}\n\n支持的格式：\n- PDF：.pdf\n- Word：.docx\n- 文本：.txt, .md\n\n请将简历文件放入该目录"
    
    result = f"✅ 找到 {len(resume_files)} 个简历文件\n\n"
    result += f"**目录**：{directory}\n\n"
    
    for idx, file_info in enumerate(resume_files, 1):
        # 格式化文件大小
        size_kb = file_info['size'] / 1024
        if size_kb < 1024:
            size_str = f"{size_kb:.1f} KB"
        else:
            size_str = f"{size_kb / 1024:.1f} MB"
        
        # 格式化文件类型
        type_map = {
            '.pdf': 'PDF',
            '.docx': 'Word',
            '.txt': 'TXT',
            '.md': 'Markdown'
        }
        
        result += f"{idx}. **{file_info['name']}**\n"
        result += f"   - 类型：{type_map.get(file_info['type'], file_info['type'])}\n"
        result += f"   - 大小：{size_str}\n"
        result += f"   - 路径：{file_info['path']}\n"
        result += f"   - 命令：`read_resume_file('{file_info['path']}')`\n\n"
    
    result += "---\n\n**使用示例**：\n"
    result += "请帮我分析简历：`read_resume_file('路径/文件名.pdf')`\n\n"
    result += "或者直接说：\"请分析我的简历\"，我会列出所有文件供你选择。"
    
    return result
